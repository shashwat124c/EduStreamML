from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set background fill colour of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """Set borders on a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = qn(f'w:{edge}')
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), kwargs.get('val', 'single'))
        element.set(qn('w:sz'), kwargs.get('sz', '6'))
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), kwargs.get('color', 'CCCCCC'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_run_with_font(para, text, bold=False, italic=False,
                      size=11, color=None, font_name='Calibri'):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def heading_style(doc, text, level):
    """Add a styled heading with blue accent colour."""
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = 'Calibri'
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)   # deep blue
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)   # medium blue
        elif level == 3:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x37, 0x51, 0x8F)
    return p

def add_paragraph(doc, text, size=11, bold=False, italic=False,
                  color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    add_run_with_font(p, text, bold=bold, italic=italic,
                      size=size, color=color)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        add_run_with_font(p, bold_prefix, bold=True, size=11)
        add_run_with_font(p, text, size=11)
    else:
        add_run_with_font(p, text, size=11)

def horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'BFDBFE')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def add_section_box(doc, title, content_lines):
    """Adds a visually distinct labelled box paragraph."""
    heading_style(doc, title, 2)
    for line in content_lines:
        add_paragraph(doc, line, size=11)

# ─────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────

def main():
    doc = Document()

    # ---------- Page margins ----------
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════
    # COVER / TITLE BLOCK
    # ══════════════════════════════════════════
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(24)
    title_para.paragraph_format.space_after  = Pt(6)

    title_run = title_para.add_run('EduStreamAI')
    title_run.bold = True
    title_run.font.size = Pt(32)
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_para.paragraph_format.space_after = Pt(4)
    sub_run = sub_para.add_run('An Intelligent, Path-Based Learning Management System')
    sub_run.italic = True
    sub_run.font.size = Pt(14)
    sub_run.font.name = 'Calibri'
    sub_run.font.color.rgb = RGBColor(0x44, 0x63, 0xA0)

    doc.add_paragraph()

    # Accent bar (1-cell table used as a coloured rule)
    bar = doc.add_table(rows=1, cols=1)
    bar.rows[0].height = Pt(6)
    bar_cell = bar.rows[0].cells[0]
    set_cell_bg(bar_cell, '1E40AF')
    bar_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    bar_cell.paragraphs[0].paragraph_format.space_before = Pt(0)

    doc.add_paragraph()

    # ══════════════════════════════════════════
    # 1. PROJECT TITLE (formal block)
    # ══════════════════════════════════════════
    heading_style(doc, '1.  Project Title', 1)
    horizontal_rule(doc)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run_with_font(p, 'Title: ', bold=True, size=12)
    add_run_with_font(p, 'EduStreamAI — Intelligent, Path-Based Learning Management System', size=12)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(12)
    add_run_with_font(p2, 'Tagline: ', bold=True, size=11)
    add_run_with_font(p2, '"Master any subject the right way — one prerequisite at a time."',
                      italic=True, size=11, color=(0x37, 0x51, 0x8F))

    # ══════════════════════════════════════════
    # 2. GROUP DETAILS
    # ══════════════════════════════════════════
    heading_style(doc, '2.  Group Details', 1)
    horizontal_rule(doc)

    p = doc.add_paragraph()
    add_run_with_font(p, 'Group Number: ', bold=True, size=11)
    add_run_with_font(p, '[Insert Group Number]', size=11)
    p.paragraph_format.space_after = Pt(8)

    # Member table
    members_table = doc.add_table(rows=5, cols=3)
    members_table.style = 'Table Grid'
    members_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    col_widths = [Inches(0.8), Inches(2.2), Inches(2.8)]
    headers = ['Sr.', 'Enrollment Number', 'Student Name']
    for i, (cell, hdr) in enumerate(zip(members_table.rows[0].cells, headers)):
        set_cell_bg(cell, '1E40AF')
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    placeholders = [
        ('1', '[Enrollment No.]', '[Student Name 1]'),
        ('2', '[Enrollment No.]', '[Student Name 2]'),
        ('3', '[Enrollment No.]', '[Student Name 3]'),
        ('4', '[Enrollment No.]', '[Student Name 4]'),
    ]
    row_colors = ['EFF6FF', 'DBEAFE']
    for r_idx, (sr, enr, name) in enumerate(placeholders):
        row = members_table.rows[r_idx + 1]
        hex_bg = row_colors[r_idx % 2]
        data = [sr, enr, name]
        for c_idx, (cell, val) in enumerate(zip(row.cells, data)):
            set_cell_bg(cell, hex_bg)
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════
    # 3. PROJECT ABSTRACT
    # ══════════════════════════════════════════
    heading_style(doc, '3.  Project Abstract', 1)
    horizontal_rule(doc)

    abstract = (
        "EduStreamAI is a modern, full-stack web application designed to help students master "
        "Computer Science and Artificial Intelligence subjects through structured dependency graphs "
        "and diagnostic knowledge checks. Traditional learning platforms deliver content linearly, "
        "regardless of whether the learner actually has the required foundational knowledge. "
        "EduStreamAI solves this gap by enforcing a prerequisite-based learning model: before a "
        "student can progress to an advanced topic, they must demonstrate understanding of its "
        "foundational concepts via an integrated MCQ quiz."
    )
    add_paragraph(doc, abstract, size=11, space_after=8)

    add_paragraph(doc,
        "When a learner fails a knowledge check, the system intelligently identifies the "
        "prerequisite gap and recommends the exact module the student should revisit — "
        "constructing a personalised learning path in real-time. The platform also monitors "
        "video-watching behaviour (excess pausing, rewinding) to detect struggle and "
        "proactively suggests an alternative video explanation.",
        size=11, space_after=8)

    heading_style(doc, 'Key Functionalities', 2)
    bullets = [
        ('Dynamic Curriculum Roadmap: ', 'Eight course domains — Programming Fundamentals, Data Structures, Algorithms, Database Internals, Artificial Intelligence, Web Development, Operating Systems, and Computer Networks — each broken into granular topics.'),
        ('MCQ Knowledge Gateways: ', 'Every topic starts with a diagnostic quiz (pass mark: 70 %) before access to the lesson video is granted.'),
        ('Intelligent Prerequisite Routing: ', 'On quiz failure, the dependency graph resolves which prerequisite topic to study and renders a curated learning-path card.'),
        ('Struggle Detection: ', 'The system tracks video pause and rewind counts; after threshold breaches, it offers an alternative explanation video.'),
        ('Progress Tracking & Auth: ', 'Google OAuth 2.0 authenticates users. Completed topic IDs are persisted in MongoDB so progress survives browser sessions.'),
    ]
    for bold_txt, rest in bullets:
        add_bullet(doc, rest, bold_prefix=bold_txt)

    heading_style(doc, 'Technologies Used', 2)
    tech_line = (
        'React.js  •  Vite  •  Tailwind CSS v4  •  Framer Motion  •  Axios  (Frontend)   |   '
        'Python  •  Flask  •  Flask-CORS  •  PyMongo  (Backend)   |   '
        'MongoDB  (Database)   |   Google OAuth 2.0  (Authentication)'
    )
    add_paragraph(doc, tech_line, size=10, italic=True,
                  color=(0x37, 0x51, 0x8F), space_after=12)

    # ══════════════════════════════════════════
    # 4. HIGH-LEVEL DESIGN
    # ══════════════════════════════════════════
    heading_style(doc, '4.  High-Level Design', 1)
    horizontal_rule(doc)

    # --- 4a System Architecture ---
    heading_style(doc, '4.1  System Architecture', 2)
    add_paragraph(doc,
        "EduStreamAI follows a classic Client-Server (SPA + REST API) architecture. "
        "The React SPA executes entirely in the browser and communicates with the Flask "
        "backend over HTTP via Axios. The backend in turn reads from and writes to a locally "
        "hosted MongoDB instance. All sensitive operations (token verification, progress "
        "updates) are proxied through the Flask layer so that no direct database access is "
        "exposed to the browser.",
        size=11, space_after=8)

    # ASCII architecture diagram as a shaded paragraph
    arch_para = doc.add_paragraph()
    arch_para.paragraph_format.left_indent  = Inches(0.4)
    arch_para.paragraph_format.space_after  = Pt(8)
    arch_run = arch_para.add_run(
        "┌───────────────────┐        HTTP / REST        ┌──────────────────────────┐\n"
        "│  React SPA        │ ─────────────────────────► │  Flask Backend (port 5000)│\n"
        "│  (Vite, port 5173)│ ◄───── JSON Responses ──── │  app.py                  │\n"
        "└───────────────────┘                            └──────────┬───────────────┘\n"
        "        │                                                   │ PyMongo\n"
        "   Google OAuth                                   ┌─────────▼──────────────┐\n"
        "   (JWT tokens)                                   │  MongoDB               │\n"
        "                                                  │  learning_path_db      │\n"
        "                                                  │  collections: modules, │\n"
        "                                                  │              users     │\n"
        "                                                  └────────────────────────┘"
    )
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(9)
    arch_run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    # --- 4b System API ---
    heading_style(doc, '4.2  System API (RESTful Endpoints)', 2)

    api_table = doc.add_table(rows=4, cols=4)
    api_table.style = 'Table Grid'
    api_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    api_col_widths = [Inches(0.7), Inches(2.0), Inches(1.5), Inches(2.0)]
    api_headers = ['Method', 'Endpoint', 'Auth Required', 'Description']

    for c, (cell, hdr) in enumerate(zip(api_table.rows[0].cells, api_headers)):
        set_cell_bg(cell, '1E40AF')
        cell.width = api_col_widths[c]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'

    api_data = [
        ('GET',  '/api/modules',       'No',   'Returns the complete curriculum — all courses, topics, prerequisites, quiz questions, and video URLs — from MongoDB.'),
        ('POST', '/api/auth/google',   'JWT',  'Verifies Google ID token, creates user profile on first login, returns user document with completed_topics array.'),
        ('POST', '/api/progress',      'JWT',  'Verifies Google token, then appends the submitted topic_id to the user\'s completed_topics set in MongoDB.'),
    ]
    api_row_colors = ['EFF6FF', 'DBEAFE']
    for r_idx, (method, endpoint, auth, desc) in enumerate(api_data):
        row = api_table.rows[r_idx + 1]
        hex_bg = api_row_colors[r_idx % 2]
        vals = [method, endpoint, auth, desc]
        for c, (cell, val) in enumerate(zip(row.cells, vals)):
            set_cell_bg(cell, hex_bg)
            cell.width = api_col_widths[c]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c < 3 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri' if c != 1 else 'Courier New'
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            if c == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- 4c Database Architecture ---
    heading_style(doc, '4.3  Database Architecture', 2)
    add_paragraph(doc,
        "EduStreamAI uses MongoDB (NoSQL document store, running on localhost:27017) with "
        "the database named learning_path_db. Below are the two primary collections:",
        size=11, space_after=6)

    heading_style(doc, 'Collection: modules', 3)
    db_bullets = [
        'topic_id (String) — Unique identifier for each topic (e.g., "prog_basics_1", "ai_2").',
        'course (String) — Name of the parent course (e.g., "Artificial Intelligence").',
        'name (String) — Human-readable topic name.',
        'difficulty (String) — "beginner" | "intermediate" | "advanced".',
        'learning_objective (String) — One-sentence outcome description.',
        'primary_video_url / fallback_video_url (String) — YouTube links for main and alternative explanations.',
        'prerequisites (Array<String>) — skill IDs or topic names this topic depends on.',
        'skills (Array<String>) — Skill tags this topic teaches (used by dependants\' prerequisites).',
        'quiz (Array<Object>) — Array of MCQ objects: { question, options[], correct_index, explanation }.',
    ]
    for b in db_bullets:
        add_bullet(doc, b)

    heading_style(doc, 'Collection: users', 3)
    user_bullets = [
        'email (String) — Google account email (unique key).',
        'name (String) — Full display name from Google profile.',
        'picture (String) — Profile picture URL.',
        'role (String) — Defaults to "student".',
        'completed_topics (Array<String>) — Set of topic_id values the user has passed and marked complete.',
    ]
    for b in user_bullets:
        add_bullet(doc, b)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- 4d Frontend Modules ---
    heading_style(doc, '4.4  Frontend Module Structure (React SPA)', 2)
    fe_bullets = [
        ('CourseLanding (route: /): ', 'Home page. Derives unique course names from the modules array, renders a responsive grid of coloured course cards with hover animations.'),
        ('TopicList (route: /course/:courseId): ', 'Filters topics by course. Displays each topic\'s name, learning objective, difficulty badge, prerequisite labels, and completion state.'),
        ('KnowledgeCheck (route: /topic/:topicId): ', 'Multi-screen component — quiz interface → results/learning-path → embedded YouTube lesson. Integrates struggle detection (pause/rewind tracking) and calls /api/progress on completion.'),
    ]
    for bold_txt, rest in fe_bullets:
        add_bullet(doc, rest, bold_prefix=bold_txt)

    # ══════════════════════════════════════════
    # 5. COMPONENTS / CONCEPTS USED
    # ══════════════════════════════════════════
    heading_style(doc, '5.  Components / Concepts Used', 1)
    horizontal_rule(doc)

    comp_data = [
        ('React.js',          'Core frontend library used to build reusable, state-driven UI components (CourseLanding, TopicList, KnowledgeCheck).'),
        ('Vite',              'Ultra-fast build tool and HMR dev-server; bundles the SPA for both development and production.'),
        ('Tailwind CSS v4',   'Utility-first CSS framework used for all styling — responsive grid layouts, hover effects, gradients, and badge components.'),
        ('Framer Motion',     'Animation library used for smooth page-transition animations (fade + blur effects via AnimatePresence).'),
        ('React Router v6',   'Client-side routing enabling SPA navigation across /, /course/:id, and /topic/:id without full page reloads.'),
        ('Axios',             'Promise-based HTTP client for API calls to the Flask backend.'),
        ('react-youtube',     'YouTube player wrapper component used to embed lesson videos with event callbacks (onReady, onStateChange).'),
        ('Lucide React',      'Icon library providing consistent iconography (BookOpen, Lock, CheckCircle2, ArrowRight, etc.).'),
        ('@react-oauth/google','Google OAuth 2.0 client library providing the GoogleOAuthProvider context and GoogleLogin button.'),
        ('Python / Flask',    'Backend web framework used to define and serve the three RESTful API endpoints.'),
        ('Flask-CORS',        'Extension that attaches CORS headers so the React dev-server (port 5173) can communicate with Flask (port 5000).'),
        ('PyMongo',           'Official MongoDB driver for Python; used to query and update the modules and users collections.'),
        ('google-auth',       'Google\'s Python library used to verify the JWT ID tokens sent from the frontend.'),
        ('MongoDB (NoSQL)',   'Document database that stores the entire curriculum (modules) and user progress (users) as flexible JSON-like documents.'),
        ('Dependency Graphs', 'Core logical concept — topics declare prerequisite skill IDs, forming a directed acyclic graph that the router traverses on quiz failure.'),
        ('JWT / OAuth 2.0',   'Industry-standard token-based authentication: the browser sends a Google-signed JWT; Flask verifies it server-side before any data mutation.'),
        ('localStorage',      'Browser storage API used to persist the logged-in user object across page refreshes without a dedicated session server.'),
        ('RESTful API Design','Stateless HTTP communication pattern with structured JSON request/response bodies and appropriate HTTP status codes.'),
        ('MCQ Quiz Engine',   'Custom quiz logic managing per-question answer tracking, scoring (≥70% = pass), and conditional routing to video or learning path.'),
        ('Struggle Detection','Heuristic using YouTube player events to count pauses (≥3) and rewinds (≥5 sec, ≥5 occurrences) as signals of learner difficulty.'),
    ]

    comp_table = doc.add_table(rows=len(comp_data) + 1, cols=2)
    comp_table.style = 'Table Grid'
    comp_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    comp_widths = [Inches(2.2), Inches(4.6)]

    for c, (cell, hdr) in enumerate(zip(comp_table.rows[0].cells, ['Component / Concept', 'Purpose in Application'])):
        set_cell_bg(cell, '1E40AF')
        cell.width = comp_widths[c]
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    for r_idx, (comp, purpose) in enumerate(comp_data):
        row = comp_table.rows[r_idx + 1]
        hex_bg = 'EFF6FF' if r_idx % 2 == 0 else 'DBEAFE'
        for c, (cell, val, is_comp) in enumerate(zip(row.cells, [comp, purpose], [True, False])):
            set_cell_bg(cell, hex_bg)
            cell.width = comp_widths[c]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.bold = is_comp
            run.font.color.rgb = (RGBColor(0x1E, 0x40, 0xAF) if is_comp
                                  else RGBColor(0x1E, 0x29, 0x3B))

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ══════════════════════════════════════════
    # 6. SCREENSHOTS AND DESCRIPTIONS
    # ══════════════════════════════════════════
    heading_style(doc, '6.  Screenshots and Descriptions', 1)
    horizontal_rule(doc)

    add_paragraph(doc,
        "Below are the key screens of the working EduStreamAI application. "
        "Please insert actual screenshots in the placeholder boxes below.",
        size=11, italic=True, color=(0x55, 0x65, 0x81), space_after=10)

    screens = [
        ('Screen 1 — Login & Authentication',
         'The landing navbar contains a Google Sign-In button (rendered by @react-oauth/google). '
         'Clicking it opens the standard Google OAuth consent popup. On success, the backend verifies the JWT, '
         'creates or retrieves the user document from MongoDB, and returns the user profile. '
         'The navbar then renders the user\'s profile picture and name in a pill-shaped badge.'),
        ('Screen 2 — Course Catalogue (Home Page)',
         'A responsive card grid shows all available course domains. Each card has a coloured gradient '
         'header, a subject-specific icon, and an "Explore" button. Cards elevate on hover. '
         'Courses available: Programming Fundamentals, Data Structures, Algorithms, Database Internals, '
         'Artificial Intelligence, Web Development, Operating Systems, and Computer Networks.'),
        ('Screen 3 — Topic List',
         'After selecting a course, the user sees a full-width gradient header with the course name and '
         'a vertical list of topic cards. Each card shows the topic name (blue), learning objective, '
         'difficulty badge, prerequisite labels (amber), and a green "Completed" badge if the user has '
         'already passed the knowledge check.'),
        ('Screen 4 — Knowledge Check / Quiz Interface',
         'A white card displays the question (with a violet icon), progress indicator ("Question X of N"), '
         'and four clickable answer options. Selecting an option highlights it in violet. '
         '"Next Question" and "Submit Answer" buttons are conditionally rendered.'),
        ('Screen 5 — Quiz Passed: Lesson Video',
         'On passing (≥70%), a success score badge and "Start Lesson Video" button appear. '
         'Clicking it renders a 16:9 embedded YouTube player. '
         'The struggle detection overlay appears after 3 pauses or 5 rewinds, offering an alternative video. '
         'A "Mark Complete →" button updates the user\'s progress via /api/progress.'),
        ('Screen 6 — Quiz Failed: Recommended Learning Path',
         'On failing, the system resolves prerequisite modules and renders an orange gradient "Start Learning" '
         'card for each recommended prerequisite. "Retry Question" and "Choose Another Topic" action buttons '
         'are also provided, plus a subtle "proceed anyway" escape link.'),
    ]

    for title, desc in screens:
        heading_style(doc, title, 2)

        # screenshot placeholder box
        ss_table = doc.add_table(rows=1, cols=1)
        ss_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        ss_cell = ss_table.rows[0].cells[0]
        ss_cell.height = Inches(2.2)
        set_cell_bg(ss_cell, 'F1F5F9')
        p = ss_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('[  Insert Screenshot Here  ]')
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
        run.italic = True
        run.font.name = 'Calibri'

        doc.add_paragraph()
        add_paragraph(doc, desc, size=11, space_after=10)

    # ══════════════════════════════════════════
    # 7. INDIVIDUAL CONTRIBUTIONS
    # ══════════════════════════════════════════
    heading_style(doc, '7.  Details of Individual Contributions', 1)
    horizontal_rule(doc)

    add_paragraph(doc,
        "The following table outlines each team member's primary area of responsibility "
        "during the development of EduStreamAI.",
        size=11, space_after=8)

    contrib_table = doc.add_table(rows=5, cols=3)
    contrib_table.style = 'Table Grid'
    contrib_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    contrib_widths = [Inches(2.2), Inches(1.8), Inches(2.8)]

    for c, (cell, hdr) in enumerate(zip(contrib_table.rows[0].cells,
                                        ['Student Name', 'Enrollment No.', 'Key Contributions'])):
        set_cell_bg(cell, '1E40AF')
        cell.width = contrib_widths[c]
        p = cell.paragraphs[0]
        run = p.add_run(hdr)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

    contrib_placeholders = [
        ('[Student Name 1]', '[Enrollment No.]',
         'Frontend SPA development — CourseLanding, TopicList, routing, Framer Motion animations.'),
        ('[Student Name 2]', '[Enrollment No.]',
         'KnowledgeCheck component — quiz engine, struggle detection, YouTube integration.'),
        ('[Student Name 3]', '[Enrollment No.]',
         'Flask backend — API endpoint design, Google OAuth verification, CORS configuration.'),
        ('[Student Name 4]', '[Enrollment No.]',
         'MongoDB schema design, seed_db.py curriculum data (8 courses, 20+ topics, quizzes).'),
    ]

    for r_idx, (name, enr, contrib) in enumerate(contrib_placeholders):
        row = contrib_table.rows[r_idx + 1]
        hex_bg = 'EFF6FF' if r_idx % 2 == 0 else 'DBEAFE'
        for c, (cell, val) in enumerate(zip(row.cells, [name, enr, contrib])):
            set_cell_bg(cell, hex_bg)
            cell.width = contrib_widths[c]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ══════════════════════════════════════════
    # 8. CONCLUSION
    # ══════════════════════════════════════════
    heading_style(doc, '8.  Conclusion', 1)
    horizontal_rule(doc)

    add_paragraph(doc,
        "EduStreamAI successfully demonstrates that an intelligent, prerequisite-aware "
        "learning management system can meaningfully improve learner outcomes compared to "
        "linear video consumption platforms. By enforcing mastery of foundational concepts "
        "before unlocking advanced topics, the platform ensures students build a cohesive, "
        "gap-free knowledge base.",
        size=11, space_after=8)

    add_paragraph(doc,
        "The project achieves a clean separation of concerns through its React SPA + "
        "Flask REST API + MongoDB architecture. The dependency-graph model powering "
        "prerequisite routing, combined with the behavioural struggle-detection heuristic, "
        "give EduStreamAI a genuine adaptive learning capability beyond basic video playlists.",
        size=11, space_after=8)

    add_paragraph(doc,
        "The integrated Google OAuth 2.0 layer ensures that personalised progress is "
        "securely associated with a verified identity and persisted across sessions — "
        "making the platform usable in real educational contexts.",
        size=11, space_after=10)

    heading_style(doc, 'Future Enhancements', 2)
    future = [
        'AI-driven quiz generation using an LLM API (e.g., Gemini) to auto-create new questions for any topic.',
        'Detailed analytics dashboard for instructors — showing class-wide struggle hotspots and topic failure rates.',
        'Peer discussion threads and collaborative notes per topic.',
        'Mobile-responsive PWA packaging for offline-capable offline video viewing.',
        'Adaptive difficulty scaling — dynamically adjusting quiz difficulty based on a student\'s historical performance.',
    ]
    for f in future:
        add_bullet(doc, f)

    doc.add_paragraph()
    add_paragraph(doc,
        "In summary, EduStreamAI is a robust proof-of-concept that bridges the gap between "
        "passive content consumption and active, verifiable learning — laying a solid "
        "foundation for a fully-fledged adaptive e-learning product.",
        size=11, italic=True, color=(0x37, 0x51, 0x8F), space_after=12)

    # ══════════════════════════════════════════
    # Save
    # ══════════════════════════════════════════
    output_path = 'EduStreamAI_Project_Report.docx'
    doc.save(output_path)
    print(f"\nReport saved successfully as: {output_path}\n")

if __name__ == '__main__':
    main()
